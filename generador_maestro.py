import os
import time
import glob
from datetime import datetime
from docx import Document
from docx2pdf import convert
import py7zr 

#! CONTRASENA FIJA
CONTRASENA_ZIP = "1234" 

#! RUTA RAIZ (Donde estan las carpetas de los Bancos)
# ".." significa una carpeta atras de donde esta el script.
RUTA_RAIZ_BUSQUEDA = os.path.dirname(os.getcwd()) 

# --- LISTA DE CLIENTES ---
CLIENTES = {
    "1": "Banco Macro",
    "2": "Banco Ciudad",
    "3": "Banco Columbia",
    "4": "Banco BST",
    "5": "Banco de Valores"
}

# Solo este cliente puede generar Informes de Pruebas Ejecutadas
CLIENTES_CON_EJECUTADO = {"Banco Macro"}

# --- FUNCION PARA ELEGIR CLIENTE ---
def seleccionar_cliente():
    """
    Muestra un menu para elegir el banco y devuelve el nombre de la carpeta.
    """
    print("\n--- Seleccione el Cliente ---")
    for key, nombre in CLIENTES.items():
        print(f"{key}. {nombre}")
    
    while True:
        opcion = input("Opcion (1-5): ")
        if opcion in CLIENTES:
            return CLIENTES[opcion]
        else:
            print("Opcion no valida. Intente de nuevo.")

# --- FUNCION INTELIGENTE DE RUTAS ---
def gestionar_ruta_destino(ast, nombre_proyecto, nombre_cliente):
    """
    1. Entra a RUTA_RAIZ / CLIENTE.
    2. Busca si existe una carpeta que empiece con el AST.
    3. Si no, crea 'AST - PROYECTO'.
    """
    # 1. Definimos la ruta del cliente (Ej: .../QA/Banco Macro)
    ruta_cliente = os.path.join(RUTA_RAIZ_BUSQUEDA, nombre_cliente)
    
    # Si la carpeta del banco no existe, la creamos para evitar errores
    if not os.path.exists(ruta_cliente):
        try:
            os.makedirs(ruta_cliente)
            print(f"[Sistema] Se creo la carpeta del cliente: {nombre_cliente}")
        except OSError:
            print(f"[Error] No se pudo crear la carpeta del cliente. Usando ruta actual.")
            return os.getcwd()

    print(f"\n[Sistema] Buscando proyecto AST {ast} en: {nombre_cliente} ...")
    
    # 2. Buscamos carpeta del proyecto DENTRO del cliente
    patron = os.path.join(ruta_cliente, f"{ast}*")
    carpetas_encontradas = glob.glob(patron)
    carpetas_encontradas = [c for c in carpetas_encontradas if os.path.isdir(c)]

    if carpetas_encontradas:
        ruta_final = carpetas_encontradas[0]
        print(f"[Sistema] Carpeta de proyecto encontrada: {os.path.basename(ruta_final)}")
    else:
        # 3. Si no existe, creamos la carpeta del proyecto
        nombre_carpeta_nueva = f"{ast} - {nombre_proyecto}"
        ruta_final = os.path.join(ruta_cliente, nombre_carpeta_nueva)
        try:
            os.makedirs(ruta_final, exist_ok=True)
            print(f"[Sistema] Carpeta nueva creada: {nombre_carpeta_nueva}")
        except Exception as e:
            print(f"[Error] Fallo al crear carpeta. Detalle: {e}")
            ruta_final = os.getcwd()

    return ruta_final

# --- LIMPIEZA DE PROCESOS ---
def limpiar_procesos():
    print("\n[Sistema] Verificando procesos de Office...")
    os.system("taskkill /f /im WINWORD.EXE >nul 2>&1")
    os.system("taskkill /f /im EXCEL.EXE >nul 2>&1")
    time.sleep(2)

# --- CONVERSION SEGURA ---
def convertir_a_pdf_seguro(docx_path, pdf_path):
    print("Convirtiendo a PDF...")
    intentos = 3
    for i in range(intentos):
        try:
            convert(docx_path, pdf_path)
            print(f"Informe (PDF) generado exitosamente.")
            return True 
        except Exception as e:
            if i < intentos - 1:
                print(f"   > Word ocupado. Reintentando en 5s... ({i+1}/{intentos})")
                time.sleep(5) 
            else:
                raise e

# --- Validaciones ---
def pedir_numero_ast():
    while True:
        n = input("Ingresa el numero de 6 digitos para AST: ")
        if n.isdigit() and len(n) == 6: return n
        print("Error: Debe ser un numero de 6 digitos.")

def pedir_texto_no_vacio(prompt):
    while True:
        t = input(prompt)
        if t.strip(): return t.replace("/", "-").replace("\\", "-").replace(":", "")
        print("Error: No puede estar vacio.")

def pedir_texto(prompt):
    return input(prompt)

# --- Reemplazo ---
def reemplazar_marcadores(doc, mapeo):
    for p in doc.paragraphs:
        for run in p.runs:
            for k, v in mapeo.items():
                if k in run.text: run.text = run.text.replace(k, v)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        for k, v in mapeo.items():
                            if k in run.text: run.text = run.text.replace(k, v)

# ==========================================
# LOGICA DE LOS INFORMES
# ==========================================

def generar_informe_c01():
    limpiar_procesos() 
    print("\n--- Generando Informe de QA de C01 ---")
    
    # 1. Elegir Cliente
    cliente = seleccionar_cliente()
    
    # 2. Datos del Proyecto
    numero_ast = pedir_numero_ast()
    proyecto = pedir_texto_no_vacio("Ingresa el nombre del Proyecto (variable para carpeta): ")
    detalle_archivo = pedir_texto_no_vacio("Ingresa 'Detalle del archivo verificado': ")
    
    # 3. Definir Ruta
    ruta_destino = gestionar_ruta_destino(numero_ast, proyecto, cliente)

    fecha_act = datetime.now(); fecha_doc = fecha_act.strftime("%d/%m/%Y")
    fecha_nom = fecha_act.strftime("%Y%m%d"); autor = "Leandro Diaz"
    
    mapeo = {
        "{{FECHA_INFORME}}": fecha_doc, "{{PROYECTO}}": proyecto,
        "{{AUTOR}}": autor, "{{NOMBRE_ARCHIVO_VERIFICADO}}": detalle_archivo,
    }
    
    nombre_base = f"FQA-102 -Informe de QA de C01 - AST {numero_ast} - {fecha_nom} (UI)"
    ruta_docx = os.path.join(ruta_destino, nombre_base + ".docx")
    ruta_pdf = os.path.join(ruta_destino, nombre_base + ".pdf")
    ruta_7z = os.path.join(ruta_destino, nombre_base + ".7z")

    try:
        doc = Document("plantilla_informe_qa.docx")
        reemplazar_marcadores(doc, mapeo)
        doc.save(ruta_docx)
        print(f"\nInforme (DOCX) guardado en: {ruta_docx}")
        
        time.sleep(2)
        convertir_a_pdf_seguro(ruta_docx, ruta_pdf)

        print("Creando 7z encriptado...")
        with py7zr.SevenZipFile(ruta_7z, 'w', password=CONTRASENA_ZIP) as z:
            z.write(ruta_pdf, arcname=os.path.basename(ruta_pdf))
        print(f"\n[EXITO] Archivos listos en carpeta: {cliente}")
        
    except Exception as e: print(f"[ERROR] {e}")


def generar_informe_codigo():
    limpiar_procesos()
    print("\n--- Generando Informe de QA de Codigo ---")

    # 1. Elegir Cliente
    cliente = seleccionar_cliente()

    numero_ast = pedir_numero_ast()
    proyecto = pedir_texto_no_vacio("Ingresa el nombre del Proyecto (variable para carpeta): ")

    comps = [] 
    n = 1
    while True:
        if n == 1: c = pedir_texto_no_vacio(f"Componente {n}: ")
        else:
            c = pedir_texto(f"Componente {n} (Enter para terminar): ")
            if not c.strip(): break
        comps.append(c); n += 1

    link_sonar = pedir_texto_no_vacio("Link SonarQube: ")
    
    # 3. Definir Ruta
    ruta_destino = gestionar_ruta_destino(numero_ast, proyecto, cliente)

    fecha_act = datetime.now(); fecha_doc = fecha_act.strftime("%d/%m/%Y")
    fecha_nom = fecha_act.strftime("%Y%m%d")
    
    lista_fmt = "\n".join([f"- {c}" for c in comps]) if comps else "N/A"
    
    mapeo = {
        "{{Fecha}}": fecha_doc, "{{ Proyecto }}": proyecto,
        "{{LISTA_COMPONENTES}}": lista_fmt, "{{SONAR}}": link_sonar,
        "{{Link_SonarQube}}": "", "{{Link_SonarCube}}": "",
        "{{Componente_1}}": "", "{{Componente_2}}": "",
        "{{Componente_3}}": "", "{{Componente_4}}": "",
    }

    nombre_base = f"FQA-100 -Informe de QA Codigo - AST {numero_ast} - {fecha_nom} (UI)"
    ruta_docx = os.path.join(ruta_destino, nombre_base + ".docx")
    ruta_pdf = os.path.join(ruta_destino, nombre_base + ".pdf")
    ruta_7z = os.path.join(ruta_destino, nombre_base + ".7z")
    
    try:
        doc = Document("plantilla_informe_qa_codigo.docx")
        reemplazar_marcadores(doc, mapeo)
        doc.save(ruta_docx)
        print(f"\nInforme (DOCX) guardado en: {ruta_docx}")
        
        time.sleep(2)
        convertir_a_pdf_seguro(ruta_docx, ruta_pdf)

        print("Creando 7z encriptado...")
        with py7zr.SevenZipFile(ruta_7z, 'w', password=CONTRASENA_ZIP) as z:
            z.write(ruta_pdf, arcname=os.path.basename(ruta_pdf))
        print(f"\n[EXITO] Archivos listos en carpeta: {cliente}")

    except Exception as e: print(f"[ERROR] {e}")


def generar_informe_ejecutado():
    limpiar_procesos()
    print("\n--- Generando Informe de Pruebas Ejecutadas ---")

    # 1. Elegir Cliente
    cliente = seleccionar_cliente()
    if cliente not in CLIENTES_CON_EJECUTADO:
        print(f"\n[Aviso] {cliente} no solicita Informes de Pruebas Ejecutadas. Operacion cancelada.")
        return

    numero_ast = pedir_numero_ast()
    proyecto = pedir_texto_no_vacio("Ingresa el nombre del Proyecto (variable para carpeta): ")

    # 3. Definir Ruta
    ruta_destino = gestionar_ruta_destino(numero_ast, proyecto, cliente)

    fecha_act = datetime.now(); fecha_doc = fecha_act.strftime("%d/%m/%Y")
    fecha_nom = fecha_act.strftime("%Y%m%d"); autor = "Leandro Diaz"

    mapeo = {
        "{{FECHA_INFORME}}": fecha_doc, "{{PROYECTO}}": proyecto, "{{AUTOR}}": autor,
    }

    nombre_base = f"FQA-104 - Informe de Revision de Pruebas Ejecutadas - AST {numero_ast} - {fecha_nom} (UI)" 
    ruta_docx = os.path.join(ruta_destino, nombre_base + ".docx")
    ruta_pdf = os.path.join(ruta_destino, nombre_base + ".pdf")
    ruta_7z = os.path.join(ruta_destino, nombre_base + ".7z")

    try:
        doc = Document("plantilla_informe_qa_Ejecutado.docx")
        reemplazar_marcadores(doc, mapeo)
        doc.save(ruta_docx)
        print(f"\nInforme (DOCX) guardado en: {ruta_docx}")
        
        time.sleep(2)
        convertir_a_pdf_seguro(ruta_docx, ruta_pdf)

        print("Creando 7z encriptado...")
        with py7zr.SevenZipFile(ruta_7z, 'w', password=CONTRASENA_ZIP) as z:
            z.write(ruta_pdf, arcname=os.path.basename(ruta_pdf))
        print(f"\n[EXITO] Archivos listos en carpeta: {cliente}")
    
    except Exception as e: print(f"[ERROR] {e}")

# --- MENU PRINCIPAL ---
def main():
    while True:
        print("\n--- Generador de Informes de QA ---")
        print("1. Informe de QA de C01")
        print("2. Informe de QA de Codigo")
        print("3. Informe de Pruebas Ejecutadas")
        print("4. Salir")
        
        op = input("Selecciona (1-4): ")
        
        if op == '1': generar_informe_c01()
        elif op == '2': generar_informe_codigo()
        elif op == '3': generar_informe_ejecutado()
        elif op == '4': break
        else: print("Opcion no valida.")

if __name__ == "__main__":
    main()